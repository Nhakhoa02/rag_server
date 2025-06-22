import PyPDF2
import pandas as pd
from PIL import Image
import pytesseract
import cv2
import numpy as np
from docx import Document
from openpyxl import load_workbook
import io
import structlog
from typing import List, Dict, Any, Optional
from pathlib import Path
import aiofiles
import asyncio

logger = structlog.get_logger()

class DocumentProcessor:
    """Process various document types and extract text"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._process_pdf,
            'docx': self._process_docx,
            'xlsx': self._process_excel,
            'csv': self._process_csv,
            'txt': self._process_text,
            'png': self._process_image,
            'jpg': self._process_image,
            'jpeg': self._process_image
        }
    
    async def process_document(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process document and extract text content"""
        try:
            if file_type not in self.supported_formats:
                raise ValueError(f"Unsupported file type: {file_type}")
            
            processor = self.supported_formats[file_type]
            content = await processor(file_path)
            
            logger.info("Document processed successfully", 
                       file_path=file_path, 
                       file_type=file_type,
                       content_length=len(content.get('text', '')))
            
            return content
            
        except Exception as e:
            logger.error("Document processing failed", 
                        file_path=file_path, 
                        file_type=file_type,
                        error=str(e))
            raise
    
    async def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF files"""
        text_content = []
        metadata = {}
        
        async with aiofiles.open(file_path, 'rb') as file:
            content = await file.read()
            
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
        
        # Extract metadata
        if pdf_reader.metadata:
            metadata = {
                'title': pdf_reader.metadata.get('/Title', ''),
                'author': pdf_reader.metadata.get('/Author', ''),
                'subject': pdf_reader.metadata.get('/Subject', ''),
                'creator': pdf_reader.metadata.get('/Creator', ''),
                'pages': len(pdf_reader.pages)
            }
        
        # Extract text from each page
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                text = page.extract_text()
                if text.strip():
                    text_content.append({
                        'page': page_num + 1,
                        'text': text.strip()
                    })
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}", error=str(e))
        
        return {
            'text': '\n\n'.join([page['text'] for page in text_content]),
            'metadata': metadata,
            'pages': text_content
        }
    
    async def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX files"""
        async with aiofiles.open(file_path, 'rb') as file:
            content = await file.read()
        
        doc = Document(io.BytesIO(content))
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text.strip())
        
        # Extract text from tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'text': '\n\n'.join(paragraphs),
            'metadata': {
                'paragraphs': len(paragraphs),
                'tables': len(tables)
            },
            'tables': tables
        }
    
    async def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel files"""
        workbook = load_workbook(file_path, data_only=True)
        
        all_data = []
        sheet_names = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = []
            
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    sheet_data.append([str(cell) if cell is not None else '' for cell in row])
            
            if sheet_data:
                all_data.extend(sheet_data)
                sheet_names.append(sheet_name)
        
        # Convert to text representation
        text_content = []
        for sheet_name in sheet_names:
            text_content.append(f"Sheet: {sheet_name}")
            sheet = workbook[sheet_name]
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    text_content.append('\t'.join([str(cell) if cell is not None else '' for cell in row]))
            text_content.append('')
        
        return {
            'text': '\n'.join(text_content),
            'metadata': {
                'sheets': sheet_names,
                'total_rows': len(all_data)
            },
            'data': all_data
        }
    
    async def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV files"""
        df = pd.read_csv(file_path)
        
        # Convert to text representation
        text_content = []
        text_content.append('\t'.join(df.columns.tolist()))
        for _, row in df.iterrows():
            text_content.append('\t'.join([str(val) for val in row.values]))
        
        return {
            'text': '\n'.join(text_content),
            'metadata': {
                'columns': df.columns.tolist(),
                'rows': len(df),
                'columns_count': len(df.columns)
            },
            'data': df.to_dict('records')
        }
    
    async def _process_text(self, file_path: str) -> Dict[str, Any]:
        """Process text files"""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            content = await file.read()
        
        return {
            'text': content,
            'metadata': {
                'lines': len(content.split('\n')),
                'characters': len(content)
            }
        }
    
    async def _process_image(self, file_path: str) -> Dict[str, Any]:
        """Process images with OCR"""
        # Read image
        image = cv2.imread(file_path)
        if image is None:
            raise ValueError("Could not read image file")
        
        # Preprocess image for better OCR
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply thresholding to get better text recognition
        _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Perform OCR
        try:
            text = pytesseract.image_to_string(thresh)
        except Exception as e:
            logger.warning("OCR failed, trying with original image", error=str(e))
            text = pytesseract.image_to_string(gray)
        
        # Get image metadata
        height, width = image.shape[:2]
        
        return {
            'text': text.strip(),
            'metadata': {
                'width': width,
                'height': height,
                'channels': image.shape[2] if len(image.shape) > 2 else 1
            }
        }
    
    def chunk_text(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """Split text into overlapping chunks"""
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings
                for i in range(end, max(start + chunk_size - 100, start), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
            if start >= len(text):
                break
        
        return chunks 